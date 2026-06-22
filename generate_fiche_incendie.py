# -*- coding: utf-8 -*-
"""Fiche programme ANDRAGOPS — Formation Incendie EPI (manipulation extincteurs + RIA),
format session 1h30, pour COUSIN GROUP. Document premium."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Palette incendie premium
ROUGE = RGBColor(0xC0, 0x39, 0x2B)      # rouge feu
ANTHRA = RGBColor(0x26, 0x32, 0x3B)     # anthracite
ORANGE = RGBColor(0xE6, 0x7E, 0x22)     # accent
GRIS = RGBColor(0x3A, 0x3A, 0x3A)
GRIS_CLAIR = "F2F3F4"
ROUGE_HEX = "C0392B"
ANTHRA_HEX = "263238"
ORANGE_HEX = "E67E22"
BLANC = RGBColor(0xFF, 0xFF, 0xFF)


def set_base(doc):
    s = doc.styles['Normal']; s.font.name = 'Calibri'; s.font.size = Pt(10.5); s.font.color.rgb = GRIS
    for sec in doc.sections:
        sec.top_margin = Cm(1.2); sec.bottom_margin = Cm(1.2)
        sec.left_margin = Cm(1.6); sec.right_margin = Cm(1.6)


def shade(cell, hexc):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), hexc)
    tcPr.append(shd)


def no_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement(f'w:{edge}'); e.set(qn('w:val'), 'none')
        borders.append(e)
    tblPr.append(borders)


def cell_text(cell, text, size=10.5, bold=False, color=GRIS, align=None, italic=False):
    cell.vertical_alignment = 1  # center
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    r = p.add_run(text); r.font.size = Pt(size); r.font.bold = bold
    r.font.color.rgb = color; r.font.italic = italic
    return r


def set_widths(table, widths):
    for row in table.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = Inches(w)


def section_title(doc, text):
    t = doc.add_table(rows=1, cols=1); no_borders(t)
    c = t.rows[0].cells[0]; shade(c, ANTHRA_HEX)
    c.width = Inches(7.0)
    p = c.paragraphs[0]
    r = p.add_run('  ' + text); r.font.size = Pt(12); r.font.bold = True; r.font.color.rgb = BLANC
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(2)


def bullet(doc, text, color_marker=ROUGE):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4); p.paragraph_format.space_after = Pt(3)
    rm = p.add_run('▪ '); rm.font.color.rgb = color_marker; rm.font.bold = True
    for i, part in enumerate(text.split('**')):
        r = p.add_run(part); r.font.size = Pt(10.5)
        if i % 2 == 1:
            r.font.bold = True
    return p


def spacer(doc, pt=4):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(pt)


doc = Document()
set_base(doc)

# ============ BANNIÈRE EN-TÊTE ============
hdr = doc.add_table(rows=1, cols=2); no_borders(hdr); set_widths(hdr, [3.7, 3.3])
left = hdr.rows[0].cells[0]; shade(left, ROUGE_HEX)
right = hdr.rows[0].cells[1]; shade(right, ROUGE_HEX)
pl = left.paragraphs[0]
r = pl.add_run('ANDRAGOPS ACADÉMIE'); r.font.size = Pt(16); r.font.bold = True; r.font.color.rgb = BLANC
pl2 = left.add_paragraph()
r = pl2.add_run('Santé, sécurité au travail & secours d’urgence'); r.font.size = Pt(8.5); r.font.color.rgb = BLANC
pr = right.paragraphs[0]; pr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = pr.add_run('FICHE PROGRAMME'); r.font.size = Pt(13); r.font.bold = True; r.font.color.rgb = BLANC
pr2 = right.add_paragraph(); pr2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = pr2.add_run('Réf. INC-EPI-1H30 · v1 · 06/2026'); r.font.size = Pt(8.5); r.font.color.rgb = BLANC
spacer(doc, 4)

# ============ TITRE ============
p = doc.add_paragraph()
r = p.add_run('Formation Incendie — Équipier de Première Intervention'); r.font.size = Pt(17)
r.font.bold = True; r.font.color.rgb = ANTHRA
p2 = doc.add_paragraph()
r = p2.add_run('Manipulation des extincteurs & RIA — exercices sur feux réels écologiques')
r.font.size = Pt(11.5); r.font.color.rgb = ROUGE; r.font.italic = True
p2.paragraph_format.space_after = Pt(6)

# ============ EN BREF (bandeau infos) ============
infos = [
    ("Public", "Tout personnel désigné (équipiers de 1ʳᵉ intervention)"),
    ("Prérequis", "Aucun"),
    ("Durée", "Session de 1 h 30 (format intra)"),
    ("Effectif", "8 participants max. / session — pour que chacun manipule"),
    ("Lieu", "Intra, sur votre site (Wervicq-Sud)"),
    ("Sanction", "Attestation individuelle + signature du registre de sécurité"),
    ("Encadrement", "Formateur spécialisé en sécurité incendie"),
    ("Accessibilité", "Référent handicap dédié — adaptations possibles"),
]
t = doc.add_table(rows=0, cols=2); no_borders(t); set_widths(t, [1.7, 5.3])
for i, (k, v) in enumerate(infos):
    row = t.add_row()
    kc, vc = row.cells
    shade(kc, ROUGE_HEX); shade(vc, GRIS_CLAIR if i % 2 == 0 else "FFFFFF")
    cell_text(kc, k, size=9.5, bold=True, color=BLANC)
    cell_text(vc, v, size=9.5)
set_widths(t, [1.7, 5.3])
spacer(doc, 6)

# ============ OBJECTIFS ============
section_title(doc, "OBJECTIFS PÉDAGOGIQUES")
para = doc.add_paragraph()
r = para.add_run("À l’issue de la session, chaque participant sera capable de :")
r.font.size = Pt(10.5)
bullet(doc, "**Identifier** un départ de feu et en comprendre les mécanismes (triangle du feu, classes de feux).")
bullet(doc, "**Donner l’alerte** et appliquer la conduite à tenir / la consigne d’évacuation de l’établissement.")
bullet(doc, "**Choisir** l’agent extincteur adapté à la classe de feu rencontrée.")
bullet(doc, "**Utiliser** un extincteur et un Robinet d’Incendie Armé (RIA) en sécurité, sur flammes réelles.")
spacer(doc, 4)

# ============ CADRE REGLEMENTAIRE ============
section_title(doc, "CADRE RÉGLEMENTAIRE")
rt = doc.add_table(rows=1, cols=1); no_borders(rt)
c = rt.rows[0].cells[0]; shade(c, GRIS_CLAIR); c.width = Inches(7.0)
cp = c.paragraphs[0]
txts = [
    ("Code du travail, art. R.4227-28 & R.4227-39 : ", "l’employeur doit doter l’établissement de moyens de lutte contre l’incendie et organiser, au moins tous les 6 mois, des exercices et essais périodiques."),
    ("Le personnel doit être entraîné", " à manipuler les moyens de premiers secours (extincteurs, RIA) et à mettre en œuvre les consignes de sécurité."),
]
for i, (b, n) in enumerate(txts):
    pp = c.add_paragraph() if i else cp
    pp.paragraph_format.space_after = Pt(2); pp.paragraph_format.left_indent = Cm(0.2)
    rb = pp.add_run(b); rb.font.bold = True; rb.font.size = Pt(9.5); rb.font.color.rgb = ANTHRA
    rn = pp.add_run(n); rn.font.size = Pt(9.5)
spacer(doc, 6)

# ============ DEROULE MINUTE ============
section_title(doc, "DÉROULÉ DE LA SESSION (1 H 30)")
d = doc.add_table(rows=1, cols=4); no_borders(d); set_widths(d, [0.9, 1.9, 3.4, 0.8])
hdrs = ["Durée", "Séquence", "Contenu", "Part.\npratique"]
for i, h in enumerate(hdrs):
    cc = d.rows[0].cells[i]; shade(cc, ANTHRA_HEX)
    cell_text(cc, h.replace('\n', ' '), size=9, bold=True, color=BLANC, align=WD_ALIGN_PARAGRAPH.CENTER)
rows = [
    ("10 min", "Accueil & cadrage", "Émargement, tour de table, objectifs, rappel des consignes et plans de sécurité du site.", "—"),
    ("25 min", "Théorie du feu", "Triangle du feu • classes de feux et pictogrammes • agents et procédés d’extinction • alerte, conduite à tenir et évacuation.", "—"),
    ("45 min", "Pratique sur feux réels", "Extincteurs (eau pulvérisée, poudre, CO₂) et RIA • extinction réelle sur générateur écologique : chaque participant manipule.", "★★★"),
    ("10 min", "Débriefing & validation", "Reprise individualisée des gestes • évaluation comportementale • signature du registre • remise des attestations.", "—"),
]
for i, row in enumerate(rows):
    rr = d.add_row()
    for j, val in enumerate(row):
        cc = rr.cells[j]
        shade(cc, "FFFFFF" if i % 2 == 0 else GRIS_CLAIR)
        al = WD_ALIGN_PARAGRAPH.CENTER if j in (0, 3) else None
        col = ROUGE if j == 3 else GRIS
        cell_text(cc, val, size=9, color=col, align=al, bold=(j == 1))
set_widths(d, [0.9, 1.9, 3.4, 0.8])
spacer(doc, 3)
note = doc.add_paragraph()
r = note.add_run("Pédagogie active — 80 % de pratique. Méthode C.A.D.R : Cadrage · Animation · Débriefing · Rapport.")
r.font.size = Pt(9); r.font.italic = True; r.font.color.rgb = ORANGE
spacer(doc, 6)

# ============ CONTENU DETAILLE (2 colonnes) ============
section_title(doc, "CONTENU DÉTAILLÉ")
cc = doc.add_table(rows=1, cols=2); no_borders(cc); set_widths(cc, [3.5, 3.5])
th = cc.rows[0].cells[0]; pr = cc.rows[0].cells[1]
# Théorie
ph = th.paragraphs[0]; r = ph.add_run('Apports théoriques'); r.font.bold = True; r.font.color.rgb = ROUGE; r.font.size = Pt(10.5)
for line in ["Consignes et plans de sécurité en place",
             "Triangle du feu : combustible, comburant, énergie",
             "Classes de feux et reconnaissance des pictogrammes",
             "Agents extincteurs et procédés d’extinction",
             "Moyens d’extinction : extincteurs & RIA"]:
    pp = th.add_paragraph(); pp.paragraph_format.space_after = Pt(2)
    rb = pp.add_run('• '); rb.font.color.rgb = ROUGE
    rr = pp.add_run(line); rr.font.size = Pt(9.5)
# Pratique
ph = pr.paragraphs[0]; r = ph.add_run('Mise en pratique'); r.font.bold = True; r.font.color.rgb = ROUGE; r.font.size = Pt(10.5)
for line in ["Présentation des extincteurs du site (eau, poudre, CO₂)",
             "Présentation et utilisation du RIA",
             "Exercices d’extinction réelle, encadrés individuellement",
             "Générateur de flammes écologique « feux propres »",
             "Sans fumée — utilisable en intérieur, protège sols & environnement"]:
    pp = pr.add_paragraph(); pp.paragraph_format.space_after = Pt(2)
    rb = pp.add_run('• '); rb.font.color.rgb = ORANGE
    rr = pp.add_run(line); rr.font.size = Pt(9.5)
set_widths(cc, [3.5, 3.5])
spacer(doc, 6)

# ============ PAGE 2 ============
doc.add_page_break()

# Moyens & évaluation
section_title(doc, "MODALITÉS, MOYENS & ÉVALUATION")
me = doc.add_table(rows=1, cols=2); no_borders(me); set_widths(me, [3.5, 3.5])
a = me.rows[0].cells[0]; b = me.rows[0].cells[1]
pa = a.paragraphs[0]; r = pa.add_run('Moyens pédagogiques & matériels'); r.font.bold = True; r.font.color.rgb = ANTHRA; r.font.size = Pt(10.5)
for line in ["Formateur spécialisé en sécurité incendie",
             "Supports visuels et démonstration",
             "Extincteurs pédagogiques (eau, poudre, CO₂)",
             "Robinet d’Incendie Armé (selon configuration)",
             "Bac à feu / générateur de flammes écologique"]:
    pp = a.add_paragraph(); pp.paragraph_format.space_after = Pt(2)
    rb = pp.add_run('• '); rb.font.color.rgb = ANTHRA
    rr = pp.add_run(line); rr.font.size = Pt(9.5)
pb = b.paragraphs[0]; r = pb.add_run('Évaluation & validation'); r.font.bold = True; r.font.color.rgb = ANTHRA; r.font.size = Pt(10.5)
for line in ["Contrôle continu du comportement en situation",
             "Validation des gestes par le formateur",
             "Émargement par demi-session",
             "Attestation de fin de formation individuelle",
             "Mémo remis à chaque participant + rapport employeur"]:
    pp = b.add_paragraph(); pp.paragraph_format.space_after = Pt(2)
    rb = pp.add_run('• '); rb.font.color.rgb = ANTHRA
    rr = pp.add_run(line); rr.font.size = Pt(9.5)
set_widths(me, [3.5, 3.5])
spacer(doc, 8)

# Organisation multi-sessions
section_title(doc, "ORGANISATION POUR PLUSIEURS GROUPES — JOURNÉE TYPE")
para = doc.add_paragraph()
r = para.add_run("Format pensé pour mettre rapidement vos équipes en conformité, sans bloquer la production. "
                 "Sessions enchaînées de 1 h 30, organisables par entité (Cousin Trestec / Cousin Composites).")
r.font.size = Pt(10)
spacer(doc, 2)
jt = doc.add_table(rows=1, cols=3); no_borders(jt); set_widths(jt, [2.3, 2.3, 2.4])
for i, h in enumerate(["Créneau", "Session", "Participants"]):
    cc = jt.rows[0].cells[i]; shade(cc, ROUGE_HEX)
    cell_text(cc, h, size=9.5, bold=True, color=BLANC, align=WD_ALIGN_PARAGRAPH.CENTER)
planning = [
    ("08 h 30 – 10 h 00", "Session 1", "jusqu’à 8"),
    ("10 h 15 – 11 h 45", "Session 2", "jusqu’à 8"),
    ("13 h 30 – 15 h 00", "Session 3", "jusqu’à 8"),
    ("15 h 15 – 16 h 45", "Session 4", "jusqu’à 8"),
]
for i, row in enumerate(planning):
    rr = jt.add_row()
    for j, val in enumerate(row):
        cc = rr.cells[j]; shade(cc, "FFFFFF" if i % 2 == 0 else GRIS_CLAIR)
        cell_text(cc, val, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
set_widths(jt, [2.3, 2.3, 2.4])
spacer(doc, 2)
tot = doc.add_paragraph()
r = tot.add_run("→ Jusqu’à 32 personnes formées par jour et par formateur. Nombre de demi-journées à caler selon vos effectifs et entités.")
r.font.size = Pt(9.5); r.font.bold = True; r.font.color.rgb = ROUGE
spacer(doc, 8)

# Financement
section_title(doc, "FINANCEMENT & ENGAGEMENT QUALITÉ")
bullet(doc, "**Organisme certifié Qualiopi** — prise en charge possible via votre OPCO (plan de développement des compétences).", ORANGE)
bullet(doc, "**Devis personnalisé** établi selon le nombre de sessions et d’entités, sous 48 h.", ORANGE)
bullet(doc, "**Convention de formation** conforme et rapport d’intervention remis à l’employeur.", ORANGE)
spacer(doc, 10)

# ============ PIED DE PAGE CONTACT ============
ft = doc.add_table(rows=1, cols=1); no_borders(ft)
c = ft.rows[0].cells[0]; shade(c, ANTHRA_HEX); c.width = Inches(7.0)
p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('ANDRAGOPS Académie'); r.font.size = Pt(11); r.font.bold = True; r.font.color.rgb = BLANC
p2 = c.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p2.add_run('Damien Lemort — Formateur SST INRS / Formateur de formateurs'); r.font.size = Pt(9); r.font.color.rgb = BLANC
p3 = c.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p3.add_run('07 59 73 82 72   ·   info@andragops-academie.com   ·   www.andragops-academie.fr')
r.font.size = Pt(9); r.font.color.rgb = BLANC
p4 = c.add_paragraph(); p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p4.add_run('Organisme de formation certifié Qualiopi · Document non contractuel'); r.font.size = Pt(7.5); r.font.italic = True; r.font.color.rgb = RGBColor(0xC8, 0xCD, 0xD0)

doc.save('/home/user/code-rouge/ANDRAGOPS_Fiche_Programme_Incendie_EPI_COUSIN_GROUP.docx')
print("Fiche programme incendie OK")
