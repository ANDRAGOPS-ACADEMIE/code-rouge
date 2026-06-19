# -*- coding: utf-8 -*-
"""Génère deux documents Word : dossier de recherche + préparation entretien COUSIN GROUP."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Couleurs charte (bleu/rouge sobre)
BLEU = RGBColor(0x1F, 0x3A, 0x5F)
ROUGE = RGBColor(0xB3, 0x1B, 0x1B)
GRIS = RGBColor(0x44, 0x44, 0x44)


def set_base_style(doc):
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.font.color.rgb = GRIS


def add_cover(doc, title, subtitle, meta_lines):
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.font.size = Pt(30)
    r.font.bold = True
    r.font.color.rgb = BLEU

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    r.font.size = Pt(15)
    r.font.color.rgb = ROUGE

    doc.add_paragraph()
    # ligne
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('—' * 20)
    r.font.color.rgb = BLEU

    for line in meta_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.size = Pt(11)
        r.font.color.rgb = GRIS
    doc.add_page_break()


def h1(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(12)
    r = p.add_run(text)
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = BLEU
    # bordure basse
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), 'B31B1B')
    pbdr.append(bottom)
    pPr.append(pbdr)
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(12.5)
    r.font.bold = True
    r.font.color.rgb = ROUGE
    return p


def para(doc, text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.bold = bold
    return p


def bullet(doc, text, sub=False):
    p = doc.add_paragraph(style='List Bullet 2' if sub else 'List Bullet')
    # support **gras** simple
    parts = text.split('**')
    for i, part in enumerate(parts):
        r = p.add_run(part)
        if i % 2 == 1:
            r.font.bold = True
    return p


def add_info_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for k, v in rows:
        cells = table.add_row().cells
        cells[0].width = Inches(2.3)
        cells[1].width = Inches(4.0)
        rk = cells[0].paragraphs[0].add_run(k)
        rk.font.bold = True
        rk.font.size = Pt(10.5)
        rk.font.color.rgb = BLEU
        rv = cells[1].paragraphs[0].add_run(v)
        rv.font.size = Pt(10.5)
    doc.add_paragraph()
    return table


def add_qa(doc, question, conseil):
    p = doc.add_paragraph()
    r = p.add_run('Q. ' + question)
    r.font.bold = True
    r.font.color.rgb = BLEU
    r.font.size = Pt(11)
    p2 = doc.add_paragraph()
    r2 = p2.add_run('Piste de réponse : ')
    r2.font.bold = True
    r2.font.color.rgb = ROUGE
    r3 = p2.add_run(conseil)
    p2.paragraph_format.space_after = Pt(10)


# =====================================================================
# DOCUMENT 1 : DOSSIER DE RECHERCHE
# =====================================================================
doc = Document()
set_base_style(doc)

add_cover(
    doc,
    "COUSIN GROUP",
    "Dossier de recherche — Préparation entretien RH",
    [
        "Cousin Trestec • Cousin Composites • Cousin Surgery (ex-Biotech)",
        "59117 Wervicq-Sud — France",
        "www.cousin-trestec.com • www.cousin-composites.com",
        "",
        "Document préparé le 19 juin 2026",
    ],
)

# Carte d'identité
h1(doc, "1. Carte d'identité du groupe")
para(doc, "Le COUSIN® Group est un groupe industriel familial français, héritier de la maison "
          "Cousin Frères fondée en 1848. Il rassemble aujourd'hui plusieurs sociétés spécialisées "
          "dans la transformation des fibres : cordages techniques, composites et dispositifs "
          "médicaux implantables. Le siège historique est situé à Wervicq-Sud, dans le Nord (59), "
          "à la frontière belge.")
add_info_table(doc, [
    ("Raison sociale", "COUSIN® Group (Cousin Trestec, Cousin Composites, Cousin Surgery)"),
    ("Création / origine", "1848 (Cousin Frères, par Louis Cousin) — sociétés actuelles créées en 1994-1995"),
    ("Siège", "8 rue de l'Abbé Bonpain, 59117 Wervicq-Sud, France"),
    ("Secteur", "Industrie textile technique, composites, dispositifs médicaux"),
    ("Forme juridique", "SAS (sociétés par actions simplifiées)"),
    ("Effectif", "≈ 150 collaborateurs (Cousin Trestec/Composites) ; ~75 pour Cousin Trestec seul"),
    ("Chiffre d'affaires", "≈ 13,5–14 M€ (Cousin Trestec, exercice 2024)"),
    ("Actionnariat", "Groupe Cousin / Dalle & Associés (groupe familial)"),
    ("Site web", "www.cousin-trestec.com — www.cousin-composites.com — www.cousin-group.com"),
])

# Histoire
h1(doc, "2. Histoire & héritage (175 ans de savoir-faire)")
bullet(doc, "**1848** : fondation de Cousin Frères par Louis (et Jean-Baptiste) Cousin, "
            "pour la fabrication de fil à coudre en lin et de lacets, dans la région de Comines.")
bullet(doc, "**Fin XIXᵉ (vers 1890)** : implantation d'un site à Wervicq-Sud (filature, retordage).")
bullet(doc, "**Début XXᵉ** : construction d'une usine rue de l'Abbé-Bonpain à Wervicq-Sud pour la fabrication de tresses.")
bullet(doc, "**Après 1918** : Vincent Cousin (également maire de Comines) recentre et développe le site de "
            "Wervicq-Sud, les usines de Comines ayant été détruites pendant la guerre.")
bullet(doc, "**1994** : création des sociétés actuelles Cousin Trestec et Cousin Composites.")
bullet(doc, "**1995** : François Cousin crée Cousin Biotech (dispositifs médicaux implantables).")
bullet(doc, "**2021** : Cousin Biotech devient Cousin Surgery ; réorganisation des filiales médicales en 2022.")
bullet(doc, "Le groupe revendique aujourd'hui **plus de 175 ans d'expertise** dans la transformation des fibres.")

# Les sociétés
h1(doc, "3. Les sociétés du groupe")

h2(doc, "Cousin Trestec — cordages et tresses techniques")
para(doc, "Cœur de métier historique. Spécialiste de la transformation des fibres synthétiques par "
          "torsion, câblage, tressage et imprégnation, pour produire cordages et cordes techniques.")
bullet(doc, "**Marchés sport / loisirs** : voile et nautisme, kitesurf, escalade, spéléologie, "
            "randonnée, alpinisme, parapente, parachutisme, canyoning, arboriculture.")
bullet(doc, "**Marchés industriels** : levage, manutention, mobilier, événementiel, aéronautique, "
            "offshore, militaire/défense, automobile.")
bullet(doc, "Plusieurs centaines de produits répartis sur ~7 gammes ; environ 5 % du CA investi en R&D ; "
            "équipe dédiée d'ingénieurs et techniciens à l'innovation.")
bullet(doc, "CA ≈ 13,5–14 M€ (2024), ~75 personnes.")

h2(doc, "Cousin Composites — joncs et profilés composites")
para(doc, "Conception, développement, fabrication et vente de joncs, profilés et tubes composites, "
          "ainsi que de cordes de raquettes (tennis, squash…).")
bullet(doc, "Joncs et profilés en fibre de verre, aramide ou carbone imprégnés de résine (pultrusion / extrusion).")
bullet(doc, "Savoir-faire : imprégnation, pultrusion, fils techniques, assemblages et gainage, surtressage technique.")
bullet(doc, "Applications : sport (cordes de raquettes), aiguilles de tirage, fils innovants, renforts techniques.")
bullet(doc, "Société créée en 1994 (Allée des Roses, Wervicq-Sud), code APE 2221Z.")

h2(doc, "Cousin Surgery (ex-Cousin Biotech) — dispositifs médicaux")
para(doc, "Branche médicale du groupe, créée en 1995 par François Cousin. Conçoit et fabrique des "
          "implants chirurgicaux (chirurgie viscérale, bariatrique, urologique, gynécologique, rachis).")
bullet(doc, "Site de ~1 300 m² dont ~500 m² de salle blanche.")
bullet(doc, "Depuis 2021 : Cousin Biotech → Cousin Surgery ; filiales renommées (Cousin Visceral, "
            "Cousin Endoscopy, Cousin Spine).")
bullet(doc, "Ensemble : leader français de l'implant chirurgical, ~30 M€ de CA, plus de 235 000 implants/an dans le monde.")

# Métier / savoir-faire
h1(doc, "4. Savoir-faire & technologies clés")
bullet(doc, "Transformation des fibres synthétiques : **torsion, câblage, tressage, imprégnation**.")
bullet(doc, "Maîtrise des fibres haute performance : polyester, polyamide, aramide (Kevlar®), Dyneema®/HMPE, carbone, verre.")
bullet(doc, "Traitements fonctionnels et finitions techniques (résistance, allongement contrôlé, traitements de surface).")
bullet(doc, "Composites : pultrusion, surtressage, gainage de renforts.")
bullet(doc, "Forte intégration verticale : du fil au produit fini, R&D et bureau d'études internes.")

# Valeurs
h1(doc, "5. Valeurs & culture d'entreprise")
para(doc, "Les valeurs mises en avant par la direction (Maxime Bonte / Valéry Dalle) :")
bullet(doc, "**Exigence et rigueur** — un niveau élevé de qualité.")
bullet(doc, "**Innovation et créativité** — l'innovation est présentée comme condition de survie et de croissance ; "
            "~5 % du CA en R&D, équipe d'ingénieurs/techniciens dédiée.")
bullet(doc, "**Savoir-faire et transmission** — un héritage industriel de 175 ans valorisé.")
bullet(doc, "**Esprit d'entreprise familiale** — proximité, ancrage territorial dans les Hauts-de-France.")
bullet(doc, "Collaboration étroite entre équipes commerciales et R&D pour capter et satisfaire la demande du marché.")

# Positionnement marché
h1(doc, "6. Positionnement & marché")
bullet(doc, "Référence internationale sur le marché de la **corde technique** (sport et industrie).")
bullet(doc, "Fabrication **française**, ancrage local fort (made in France, savoir-faire textile du Nord).")
bullet(doc, "Clients exigeants : sécurité (vie en jeu pour l'escalade, le travail en hauteur), défense, médical.")
bullet(doc, "Diversification : du cordage sport jusqu'à l'implant médical, via les composites techniques.")

# Gouvernance
h1(doc, "7. Gouvernance & organisation")
bullet(doc, "Cousin Trestec et Cousin Composites : filiales adossées à **Dalle & Associés / Groupe Cousin**.")
bullet(doc, "Direction connue : **Maxime Bonte** et **Valéry Dalle** (codirection Trestec/Composites) ; "
            "**Tanguy Fortin** cité comme Directeur Général de Cousin Trestec.")
bullet(doc, "Depuis le 15 septembre 2022, la société Cousin Group est nommée Présidente (en remplacement de Dalle & Associés).")
bullet(doc, "Note : les noms de dirigeants évoluent — à vérifier/confirmer le jour de l'entretien.")

# Coordonnées
h1(doc, "8. Coordonnées utiles")
add_info_table(doc, [
    ("Adresse", "8 rue de l'Abbé Bonpain, 59117 Wervicq-Sud, France"),
    ("Téléphone standard", "+33 (0)3 20 14 40 00"),
    ("Email", "contact@cousin-trestec.com"),
    ("Contact RH (entretien)", "Responsable Ressources Humaines — +33 (0)6 35 25 89 22"),
    ("Sites web", "cousin-trestec.com • cousin-composites.com • cousin-group.com"),
])

# Sources
h1(doc, "9. Sources")
for s in [
    "Site officiel du groupe : cousin-group.com",
    "societe.com / pappers.fr / manageo.fr (données légales et financières)",
    "LinkedIn Cousin Trestec & Cousin Composites",
    "modeintextile.fr — interview de Maxime Bonte, dirigeant",
    "lagazettefrance.fr — articles sur l'histoire et Cousin Surgery",
    "wervicq-sud.com — patrimoine historique ; visitwapi.be — histoire familiale",
    "Offres d'emploi : hellowork.com, indeed.fr, recrut.com",
]:
    bullet(doc, s)
para(doc, "Avertissement : informations issues de recherches publiques (juin 2026), à recouper. "
          "Les chiffres (CA, effectifs) peuvent avoir évolué.").runs[0].font.size = Pt(9)

doc.save('/home/user/code-rouge/Dossier_Recherche_COUSIN_GROUP.docx')
print("Document 1 OK")

# =====================================================================
# DOCUMENT 2 : PREPARATION ENTRETIEN
# =====================================================================
doc2 = Document()
set_base_style(doc2)

add_cover(
    doc2,
    "Préparation à l'entretien",
    "Entretien téléphonique — Responsable RH, COUSIN GROUP",
    [
        "Cousin Trestec / Cousin Composites — Wervicq-Sud (59)",
        "Contact RH : +33 (0)6 35 25 89 22",
        "",
        "Guide personnel de préparation — 19 juin 2026",
    ],
)

h1(doc2, "1. Avant l'appel — check-list")
bullet(doc2, "Relire le **dossier de recherche** sur le groupe (histoire 1848, 3 pôles, valeurs).")
bullet(doc2, "Avoir sous les yeux : l'offre de poste, mon CV, ce document.")
bullet(doc2, "Préparer un **endroit calme**, bon réseau, écouteurs, papier + stylo.")
bullet(doc2, "Vérifier le nom et la fonction exacts de l'interlocuteur RH.")
bullet(doc2, "Préparer une feuille avec : 3 atouts à placer, 3 questions à poser, mes disponibilités.")
bullet(doc2, "Appeler / décrocher avec le sourire (s'entend au téléphone), eau à portée de main.")

h1(doc2, "2. Pitch de présentation (30–60 secondes)")
para(doc2, "Structure recommandée — à adapter à votre parcours :")
bullet(doc2, "**Qui je suis** : prénom + situation actuelle (poste / formation).")
bullet(doc2, "**Mon parcours en 2 phrases** : expériences et compétences clés en lien avec le poste.")
bullet(doc2, "**Pourquoi Cousin** : intérêt pour une entreprise industrielle française, innovante, "
             "avec un savoir-faire de 175 ans dans les fibres techniques.")
bullet(doc2, "**Ce que je recherche** : une phrase sur le projet professionnel et l'envie de contribuer.")
para(doc2, "À personnaliser : notez ci-dessous votre version :")
for _ in range(3):
    doc2.add_paragraph("……………………………………………………………………………………………………")

h1(doc2, "3. Questions classiques en entretien RH (et pistes de réponse)")
add_qa(doc2, "Parlez-moi de vous / présentez votre parcours.",
       "Dérouler le pitch ci-dessus, en 1 minute max, orienté vers le poste visé.")
add_qa(doc2, "Que savez-vous de notre entreprise ?",
       "Montrer la recherche : groupe familial fondé en 1848, cordages techniques (Trestec), "
       "composites (Composites) et médical (Cousin Surgery) ; site de Wervicq-Sud ; fibres techniques ; "
       "valeurs d'exigence, qualité et innovation (~5 % du CA en R&D).")
add_qa(doc2, "Pourquoi voulez-vous travailler chez nous / pour ce poste ?",
       "Relier vos compétences aux besoins du poste + adhésion aux valeurs (industrie française, "
       "innovation, savoir-faire, qualité). Donner une raison sincère et concrète.")
add_qa(doc2, "Quelles sont vos qualités / vos défauts ?",
       "2-3 qualités illustrées d'exemples concrets ; 1 défaut honnête + comment vous le gérez.")
add_qa(doc2, "Parlez-moi d'une réussite / d'un défi que vous avez relevé.",
       "Méthode STAR : Situation, Tâche, Action, Résultat (chiffré si possible).")
add_qa(doc2, "Comment gérez-vous la pression / le travail en équipe ?",
       "Exemple vécu : organisation, priorisation, communication. Valoriser rigueur et fiabilité "
       "(important en milieu industriel où la qualité/sécurité priment).")
add_qa(doc2, "Quelles sont vos prétentions salariales ?",
       "Donner une fourchette réaliste préparée à l'avance (se renseigner sur le marché du poste/région). "
       "Rester ouvert : « selon le poste et les responsabilités »." )
add_qa(doc2, "Quelles sont vos disponibilités / votre préavis ?",
       "Réponse claire et honnête sur la date de prise de poste possible.")
add_qa(doc2, "Où vous voyez-vous dans 3 à 5 ans ?",
       "Montrer envie d'évoluer et de s'investir durablement, en cohérence avec une PME industrielle "
       "où l'on peut monter en compétences.")
add_qa(doc2, "Avez-vous des questions ?",
       "Toujours OUI — voir section 5.")

h1(doc2, "4. Spécial entretien téléphonique")
bullet(doc2, "Décrocher en se présentant clairement : « Bonjour, [Prénom Nom] à l'appareil ».")
bullet(doc2, "Parler **lentement et distinctement**, sourire, ne pas couper la parole.")
bullet(doc2, "Sans langage corporel : la **voix** porte tout — ton dynamique et posé.")
bullet(doc2, "Noter le nom de l'interlocuteur et les points clés pendant l'échange.")
bullet(doc2, "Reformuler si besoin : « Si je comprends bien… ».")
bullet(doc2, "Garder le dossier de recherche et le CV sous les yeux (avantage du téléphone !).")

h1(doc2, "5. Questions à poser au/à la RH")
bullet(doc2, "Comment décririez-vous **l'équipe** et l'environnement de travail au quotidien ?")
bullet(doc2, "Quelles seraient les **priorités** sur les premiers mois dans ce poste ?")
bullet(doc2, "Comment l'entreprise accompagne-t-elle la **montée en compétences / la formation** ?")
bullet(doc2, "Quelles sont les **perspectives d'évolution** au sein du groupe ?")
bullet(doc2, "Comment se traduit concrètement la **culture d'innovation** dont parle la direction ?")
bullet(doc2, "Quelles sont les **prochaines étapes** du processus de recrutement ?")

h1(doc2, "6. À éviter")
bullet(doc2, "Dire « je ne sais rien de l'entreprise ».")
bullet(doc2, "Critiquer un ancien employeur.")
bullet(doc2, "Répondre par monosyllabes ou, à l'inverse, monopoliser la parole.")
bullet(doc2, "Parler argent en premier / trop tôt sans qu'on le demande.")
bullet(doc2, "Être en retard, dans un lieu bruyant, ou distrait.")

h1(doc2, "7. Après l'entretien")
bullet(doc2, "Noter à chaud : points abordés, ressenti, prochaines étapes, date de relance.")
bullet(doc2, "Envoyer si pertinent un **email de remerciement** bref et personnalisé.")
bullet(doc2, "Préparer les éléments éventuellement demandés (références, documents, dispo).")

h1(doc2, "8. Notes personnelles")
for _ in range(6):
    doc2.add_paragraph("……………………………………………………………………………………………………")

doc2.save('/home/user/code-rouge/Preparation_Entretien_COUSIN_GROUP.docx')
print("Document 2 OK")
