# -*- coding: utf-8 -*-
import os, secrets, segno
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

A="Arial"; BLEU=RGBColor(0x1F,0x4E,0x79); ROUGE=RGBColor(0xC0,0x1A,0x1A); GRIS=RGBColor(0x80,0x80,0x80)
BLANC=RGBColor(0xFF,0xFF,0xFF); NOIR=RGBColor(0x22,0x22,0x22); OR=RGBColor(0xE8,0xB5,0x0A); VERT=RGBColor(0x1F,0x8A,0x4C)
FOLDER="/home/user/code-rouge/jeu_cousin"
GAME="https://andragops-academie.github.io/sang-froid-cousin/"
DASH=GAME+"tableau-mission.html"
QR=os.path.join(FOLDER,"QR_jeu.png")
ADMIN="cou-"+secrets.token_urlsafe(18).replace('-','').replace('_','')[:26]

# ---- QR ----
segno.make(GAME, error='m').save(QR, scale=11, border=2, dark="111111")

def newdoc(margins=(1.6,1.4,2,2)):
    d=Document()
    for s in d.sections:
        s.top_margin=Cm(margins[0]); s.bottom_margin=Cm(margins[1]); s.left_margin=Cm(margins[2]); s.right_margin=Cm(margins[3])
    d.styles['Normal'].font.name=A; d.styles['Normal'].font.size=Pt(11)
    return d
def run(p,t,sz=11,b=False,i=False,c=None,al=None):
    if al is not None:p.alignment=al
    r=p.add_run(t);r.font.name=A;r.font.size=Pt(sz);r.font.bold=b;r.font.italic=i
    if c:r.font.color.rgb=c
    return r
def shade(cell,hexc):
    sh=OxmlElement('w:shd');sh.set(qn('w:val'),'clear');sh.set(qn('w:fill'),hexc);cell._tc.get_or_add_tcPr().append(sh)
def borders(cell,col="BFC8D6",sz=4):
    tcPr=cell._tc.get_or_add_tcPr();b=OxmlElement('w:tcBorders')
    for e in ("top","left","bottom","right"):
        el=OxmlElement('w:'+e);el.set(qn('w:val'),'single');el.set(qn('w:sz'),str(sz));el.set(qn('w:color'),col);b.append(el)
    tcPr.append(b)

# =====================================================================
# 1) GUIDE RH / ANIMATEUR
# =====================================================================
d=newdoc((1.8,1.8,2,2))
def para(t,**k):
    p=d.add_paragraph();p.paragraph_format.space_after=Pt(5);run(p,t,**k);return p
def h(t):
    p=d.add_paragraph();p.paragraph_format.space_before=Pt(10);p.paragraph_format.space_after=Pt(3);run(p,t,12.5,True,c=BLEU);return p
def bul(t,lead=None):
    p=d.add_paragraph();p.paragraph_format.left_indent=Cm(0.6);p.paragraph_format.first_line_indent=Cm(-0.4);p.paragraph_format.space_after=Pt(3)
    run(p,"•  ",b=True,c=BLEU)
    if lead:run(p,lead,b=True)
    run(p,t);return p
p=d.add_paragraph();p.paragraph_format.space_after=Pt(6);run(p,"OPÉRATION SANG-FROID — Édition Industrie",16,True,c=BLEU)
p=d.add_paragraph();p.paragraph_format.space_after=Pt(2);run(p,"Guide RH / Animateur — ANDRAGOPS Académie pour COUSIN GROUP",11,True,c=ROUGE)
para("Jeu d'accroche interactif (type escape game), version individuelle, avec collecte centralisée des résultats et des demandes de formation. Objectif : lever le frein du volontariat — créer l'envie de se former, et fournir au service RH un plan de formation chiffré.",i=True,c=GRIS)
h("1. Le principe")
para("Chaque salarié joue seul ~8 minutes (situations de sécurité du quotidien à l'atelier + énigmes). À la fin : un badge, un mini-diagnostic personnalisé, puis une page où il COCHE les formations qui l'intéressent (SST, MAC SST, Gestes & Postures / PRAP, Incendie) + une observation. Ces choix sont transmis automatiquement au service RH.")
para("Formations ciblées : SST · MAC SST · Gestes & Postures / PRAP · Incendie (extincteurs, RIA, évacuation).")
h("2. Le jeu — À HÉBERGER")
para("Le jeu est un fichier autonome (index.html). Aucune installation pour les salariés une fois en ligne.")
bul(GAME,lead="Lien direct (à diffuser, après mise en ligne) : ")
bul("scanner QR_jeu.png (voir la fiche « LIENS_ET_QR ») — ouvre directement le jeu.",lead="QR code : ")
bul("double-clic sur index.html dans Chrome ou Edge ; garder index.html + logo.png groupés.",lead="En local (poste / tablette) : ")
h("3. Comment ça se joue")
bul("Le joueur saisit prénom + atelier/service + s'il est déjà SST (pour orienter SST initiale vs MAC SST).")
bul("17 épreuves sous chrono (situations d'atelier + énigmes). Bons réflexes = points ; erreur = temps perdu.")
bul("À la fin : badge Or/Argent/Bronze, mini-diagnostic, puis CHOIX des formations (cases pré-cochées selon le diagnostic) + observation → « transmettre au RH ».")
h("4. Le tableau de mission — tableau-mission.html (RÉSERVÉ RH / ANDRAGOPS)")
para("Ouvre tableau-mission.html dans un navigateur. Code d'accès demandé une fois :")
bul(ADMIN+"   (À GARDER CONFIDENTIEL — ne pas afficher ni imprimer publiquement)",lead="Code d'accès : ")
bul("Interrupteur ON / OFF : démarrer ou suspendre le jeu pour tout le monde (OFF = écran « fermé » côté joueur).")
bul("Onglet « Synthèse & demandes » : participation, DEMANDES DE FORMATION (nombre de volontaires par formation), liste des participants avec leurs choix + observations, analyse des lacunes.")
bul("Onglet « Radar des lacunes » : visualisation des domaines à renforcer (secours / incendie / TMS).")
bul("Bouton « Synthèse Word » : document entreprise prêt à présenter. Bouton « Export Excel » : tous les participants + leurs demandes de formation. Bouton « Réinitialiser » : remise à zéro.")
h("5. Données & cadre")
bul("Données minimales : prénom + atelier/service (facultatif). Aucune donnée sensible.")
bul("Hébergement : projet Supabase « ANDRAGOPS » (Europe, eu-west-3), table dédiée cousin_results, séparée des autres clients.")
bul("Accès admin (ON/OFF, données, exports, reset) protégé par le code ci-dessus. Le jeu côté joueur n'expose aucune donnée des autres.")
bul("Outil d'accroche pédagogique, NON certifiant — il prépare et motive, il ne remplace pas la formation.")
p=d.add_paragraph();p.paragraph_format.space_before=Pt(12);run(p,"Contact : ANDRAGOPS Académie · Damien LEMORT · 07 59 73 82 72 · info@andragops-academie.fr · www.andragops-academie.fr",9.5,i=True,c=GRIS)
d.save(os.path.join(FOLDER,"OPERATION_SANG-FROID_Cousin_guide-RH.docx"))

# =====================================================================
# 2) FICHE DE RETOUR RH
# =====================================================================
d=newdoc((1.6,1.4,2,2)); d.styles['Normal'].font.size=Pt(10.5)
def runp(p,t,sz=10.5,b=False,i=False,c=None,al=None):
    if al is not None:p.alignment=al
    r=p.add_run(t);r.font.name=A;r.font.size=Pt(sz);r.font.bold=b;r.font.italic=i
    if c:r.font.color.rgb=c
    return r
def paraf(t="",sz=10.5,b=False,i=False,c=None,al=None,after=5,before=0):
    p=d.add_paragraph();p.paragraph_format.space_after=Pt(after);p.paragraph_format.space_before=Pt(before)
    if t:runp(p,t,sz,b,i,c,al)
    return p
def cell(c,t,sz=10,b=False,col=None,bg=None,al=WD_ALIGN_PARAGRAPH.LEFT):
    borders(c)
    if bg:shade(c,bg)
    pp=c.paragraphs[0];pp.alignment=al;pp.paragraph_format.space_after=Pt(1);pp.paragraph_format.space_before=Pt(1)
    r=pp.add_run(t);r.font.name=A;r.font.size=Pt(sz);r.font.bold=b
    if col:r.font.color.rgb=col
    return c
def hf(n,t):
    p=d.add_paragraph();p.paragraph_format.space_before=Pt(11);p.paragraph_format.space_after=Pt(4)
    runp(p,n+"   ",12,True,c=ROUGE);runp(p,t,12,True,c=BLEU);return p
def line(lead="",n=1):
    for _ in range(n):
        p=d.add_paragraph();p.paragraph_format.space_after=Pt(3)
        if lead:runp(p,lead+" ",10,b=True,c=GRIS)
        runp(p,"…"*60,10,c=RGBColor(0xC8,0xCE,0xD8))
p=d.add_paragraph();runp(p,"ANDRAGOPS ACADÉMIE",13,True,c=BLEU)
p=d.add_paragraph();p.paragraph_format.space_after=Pt(2);runp(p,"OPÉRATION SANG-FROID — Fiche de retour",18,True,c=NOIR)
p=d.add_paragraph();p.paragraph_format.space_after=Pt(8);runp(p,"Votre avis sur le jeu d'accroche — il nous aide à ajuster et à construire votre plan de formation.",10.5,i=True,c=GRIS)
t=d.add_table(rows=1,cols=3);t.alignment=WD_TABLE_ALIGNMENT.CENTER
cell(t.rows[0].cells[0],"Répondant : Pauline VALLÉE (RH)",9.5)
cell(t.rows[0].cells[1],"Entreprise : Cousin Group — Wervicq-Sud",9.5)
cell(t.rows[0].cells[2],"Date : ____ / ____ / 2026",9.5)
hf("1.","Votre appréciation générale")
paraf("Marquez d'un X la case qui correspond.",9.5,i=True,c=GRIS,after=4)
crit=["Le concept (escape game d'accroche)","La qualité et le professionnalisme du jeu","La facilité de mise en place / diffusion","Le tableau de suivi RH (lisibilité, exports Word/Excel)","L'adéquation avec vos besoins"]
cols=["Excellent","Bien","Moyen","À revoir"]
t=d.add_table(rows=len(crit)+1,cols=5);t.alignment=WD_TABLE_ALIGNMENT.CENTER
cell(t.rows[0].cells[0],"Critère",9.5,True,BLANC,"1F4E79")
for j,cn in enumerate(cols):cell(t.rows[0].cells[j+1],cn,9.5,True,BLANC,"1F4E79",WD_ALIGN_PARAGRAPH.CENTER)
for i,cr in enumerate(crit):
    cell(t.rows[i+1].cells[0],cr,10)
    for j in range(4):cell(t.rows[i+1].cells[j+1],"",10,al=WD_ALIGN_PARAGRAPH.CENTER)
for row in t.rows:
    for k,w in enumerate([Cm(8.6),Cm(2.2),Cm(2.2),Cm(2.2),Cm(2.2)]):row.cells[k].width=w
hf("2.","L'objectif n°1 : lever le frein du volontariat")
paraf("Le jeu a-t-il suscité de l'intérêt pour les formations ?",10.5,True,after=2)
p=d.add_paragraph();p.paragraph_format.space_after=Pt(4)
for opt in ["☐  Oui, nettement","☐  Plutôt oui","☐  Mitigé","☐  Non"]:runp(p,opt+"        ",10.5)
paraf("Nombre de salariés ayant joué (estimation) :  ………………",10.5,after=3)
paraf("Retours spontanés entendus en interne :",10.5,True,after=2);line(n=2)
hf("3.","La suite — formations envisagées")
paraf("Cochez les formations qui vous intéressent et indiquez un nombre de salariés (même approximatif).",9.5,i=True,c=GRIS,after=4)
forms=["SST — Sauveteur Secouriste du Travail (14h)","MAC SST — recyclage SST (7h)","Gestes & Postures / PRAP — manutention & TMS (7h)","Incendie — extincteurs, RIA & évacuation","Autre (préciser) : ………………………………………"]
t=d.add_table(rows=len(forms)+1,cols=3);t.alignment=WD_TABLE_ALIGNMENT.CENTER
cell(t.rows[0].cells[0],"Formation",9.5,True,BLANC,"1F4E79")
cell(t.rows[0].cells[1],"Intéressé(e) ?",9.5,True,BLANC,"1F4E79",WD_ALIGN_PARAGRAPH.CENTER)
cell(t.rows[0].cells[2],"Nb de salariés (estim.)",9.5,True,BLANC,"1F4E79",WD_ALIGN_PARAGRAPH.CENTER)
for i,fm in enumerate(forms):
    cell(t.rows[i+1].cells[0],fm,10)
    cell(t.rows[i+1].cells[1],"☐",12,al=WD_ALIGN_PARAGRAPH.CENTER)
    cell(t.rows[i+1].cells[2],"",10,al=WD_ALIGN_PARAGRAPH.CENTER)
for row in t.rows:
    for k,w in enumerate([Cm(10.0),Cm(2.6),Cm(4.4)]):row.cells[k].width=w
paraf("Souhaitez-vous renouveler / étendre le dispositif (autres entités, nouvelles campagnes) ?",10.5,True,before=8,after=2)
p=d.add_paragraph();p.paragraph_format.space_after=Pt(4)
for opt in ["☐  Oui","☐  Peut-être","☐  Non"]:runp(p,opt+"        ",10.5)
hf("4.","Points forts / points à améliorer")
paraf("Ce qui vous a plu :",10.5,True,after=2);line(n=2)
paraf("Ce qu'on pourrait améliorer :",10.5,True,before=4,after=2);line(n=2)
hf("5.","Recommandation")
paraf("Recommanderiez-vous ANDRAGOPS à un confrère / une autre entreprise ?",10.5,True,after=2)
p=d.add_paragraph();p.paragraph_format.space_after=Pt(4)
for opt in ["☐  Oui","☐  Non"]:runp(p,opt+"        ",10.5)
paraf("Un mot que nous pourrions utiliser comme référence (facultatif) :",10.5,True,after=2);line(n=2)
paraf("Note globale du dispositif :        ……… / 10",11,True,before=4,c=BLEU)
paraf("",after=4)
t=d.add_table(rows=1,cols=2);t.alignment=WD_TABLE_ALIGNMENT.CENTER
cell(t.rows[0].cells[0],"Nom & fonction : …………………………………",9.5)
cell(t.rows[0].cells[1],"Signature / cachet :",9.5)
t.rows[0].cells[0].width=Cm(9.5);t.rows[0].cells[1].width=Cm(7.5)
p=d.add_paragraph();p.paragraph_format.space_before=Pt(12)
runp(p,"Merci pour votre retour ! — ANDRAGOPS Académie · Damien LEMORT · 07 59 73 82 72 · info@andragops-academie.fr · www.andragops-academie.fr",9,i=True,c=GRIS)
d.save(os.path.join(FOLDER,"OPERATION_SANG-FROID_Fiche-de-retour-RH.docx"))

# =====================================================================
# 3) COMMUNICATION DE LANCEMENT (e-mail + affiche)
# =====================================================================
d=newdoc((1.6,1.4,2,2))
def paral(t="",sz=11,b=False,i=False,c=None,al=None,after=6,before=0):
    p=d.add_paragraph();p.paragraph_format.space_after=Pt(after);p.paragraph_format.space_before=Pt(before)
    if t:run(p,t,sz,b,i,c,al)
    return p
def banner(lines):
    t=d.add_table(rows=1,cols=1);t.alignment=WD_TABLE_ALIGNMENT.CENTER
    c=t.rows[0].cells[0];shade(c,"1F4E79");c.paragraphs[0].text=""
    for i,(txt,sz,col,bold) in enumerate(lines):
        p=c.paragraphs[0] if i==0 else c.add_paragraph()
        p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.paragraph_format.space_after=Pt(2);p.paragraph_format.space_before=Pt(2)
        run(p,txt,sz,bold,c=col)
    return t
paral("À COPIER-COLLER — E-mail de lancement",13,True,c=ROUGE,after=8)
paral("Objet :",10,True,c=GRIS,after=1)
paral("🚨 OPÉRATION SANG-FROID — sauras-tu garder ton sang-froid ? (8 min, en ligne)",12,True,c=BLEU,after=10)
paral("Bonjour à toutes et à tous,",after=8)
paral("Et si on testait nos réflexes de sécurité… en s'amusant ?",b=True,after=6)
paral("Cousin Group lance OPÉRATION SANG-FROID, un escape game en ligne créé pour nous : 8 minutes, depuis ton poste ou ton téléphone, pour gérer une matinée mouvementée à l'atelier (un malaise, un départ de feu, une évacuation…) et décrocher ton badge.")
p=paral("👉 Pour jouer : ",b=True,after=8);run(p,GAME,11,True,c=ROUGE);run(p,"   (ou scanne le QR code ci-dessous)")
paral("Au programme :",b=True,after=3)
for tt in ["Un défi chrono, des situations concrètes, un classement entre collègues 🏆","Ton mini-diagnostic personnalisé à la fin","Et surtout : tu choisis les formations qui t'intéressent (SST, gestes & postures, incendie…) — l'entreprise s'occupe du reste."]:
    pp=d.add_paragraph();pp.paragraph_format.left_indent=Cm(0.6);pp.paragraph_format.first_line_indent=Cm(-0.4);pp.paragraph_format.space_after=Pt(3)
    run(pp,"•  ",b=True,c=BLEU);run(pp,tt)
paral("C'est ludique, non noté, ça prend 8 minutes, et seul ton prénom apparaît au classement. Prêt à relever le défi ?",after=8,before=4)
paral("À vos souris — et gardez votre sang-froid ! 😎",b=True,after=8)
paral("Le service RH",i=True,c=GRIS,after=0)
d.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
banner([("ANDRAGOPS × COUSIN GROUP",10,RGBColor(0xDC,0xE6,0xF1),False),
        ("OPÉRATION SANG-FROID",30,BLANC,True),
        ("Escape game prévention · Édition Industrie",13,RGBColor(0xDC,0xE6,0xF1),False)])
paral("",after=4)
paral("Sauras-tu sécuriser l'atelier avant la fin du chrono ?",18,True,c=BLEU,al=WD_ALIGN_PARAGRAPH.CENTER,after=4)
paral("Un malaise, un départ de feu, une évacuation… 8 minutes pour garder ton sang-froid.",12,i=True,c=GRIS,al=WD_ALIGN_PARAGRAPH.CENTER,after=10)
p=d.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
d.add_picture(QR,width=Inches(2.7));d.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
paral("📲 Scanne et joue — depuis ton poste ou ton téléphone",13,True,c=ROUGE,al=WD_ALIGN_PARAGRAPH.CENTER,after=2,before=4)
paral(GAME,11,True,c=BLEU,al=WD_ALIGN_PARAGRAPH.CENTER,after=10)
for tt in ["🏆 Un badge à décrocher + un classement entre collègues","🎯 Ton mini-diagnostic sécurité personnalisé","🎓 Tu choisis les formations qui t'intéressent"]:
    paral(tt,12,True,al=WD_ALIGN_PARAGRAPH.CENTER,after=3)
paral("Ludique · 8 minutes · non noté · seul ton prénom apparaît",10,i=True,c=GRIS,al=WD_ALIGN_PARAGRAPH.CENTER,before=8)
d.save(os.path.join(FOLDER,"Communication_Lancement_Sang-Froid_Cousin.docx"))

# =====================================================================
# 4) LIENS & QR
# =====================================================================
d=newdoc((1.8,1.8,2,2))
p=d.add_paragraph();run(p,"OPÉRATION SANG-FROID — Diffuser le jeu",16,True,c=BLEU)
p=d.add_paragraph();p.paragraph_format.space_after=Pt(8);run(p,"Édition Industrie · ANDRAGOPS Académie pour COUSIN GROUP",11,True,c=ROUGE)
p=d.add_paragraph();p.paragraph_format.space_before=Pt(6);run(p,"1. Le jeu — à diffuser aux salariés",12.5,True,c=BLEU)
p=d.add_paragraph();run(p,"Scanne le QR code ou ouvre le lien. Aucune installation. À afficher (panneau atelier, salle de pause), à mettre en signature d'e-mail ou sur l'intranet.")
p=d.add_paragraph();run(p,GAME,12,True,c=VERT)
p=d.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.paragraph_format.space_before=Pt(6)
d.add_picture(QR, width=Inches(2.4));d.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
p=d.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;run(p,"QR — Opération Sang-Froid (jeu)",9,i=True,c=GRIS)
p=d.add_paragraph();p.paragraph_format.space_before=Pt(10);run(p,"2. Le tableau de mission — RÉSERVÉ RH / ANDRAGOPS",12.5,True,c=BLEU)
p=d.add_paragraph();run(p,"Suivi en direct : participation, demandes de formation (volontaires), lacunes, exports Word/Excel, interrupteur ON/OFF.")
p=d.add_paragraph();run(p,DASH,11,True,c=BLEU)
p=d.add_paragraph();run(p,"Code d'accès : ",11,True);run(p,"confidentiel — voir le Guide RH (ne pas diffuser cette fiche avec le code)",11,True,c=ROUGE)
p=d.add_paragraph();p.paragraph_format.space_before=Pt(10);run(p,"3. Bon à savoir",12.5,True,c=BLEU)
for tt in ["Le jeu et le tableau lisent/écrivent en direct sur la plateforme (Supabase, Europe) — données séparées des autres clients (table cousin_results).",
          "Données minimales (prénom + atelier/service), non sensibles. Outil d'accroche, non certifiant.",
          "Pour fermer le jeu entre deux campagnes : interrupteur OFF dans le tableau de mission.",
          "Republier une mise à jour du jeu : me le demander (re-déploiement GitHub Pages)."]:
    pp=d.add_paragraph();pp.paragraph_format.left_indent=Cm(0.6);pp.paragraph_format.first_line_indent=Cm(-0.4);pp.paragraph_format.space_after=Pt(3)
    run(pp,"•  ",b=True,c=BLEU);run(pp,tt)
p=d.add_paragraph();p.paragraph_format.space_before=Pt(12);run(p,"Contact : ANDRAGOPS Académie · Damien LEMORT · 07 59 73 82 72 · info@andragops-academie.fr · www.andragops-academie.fr",9.5,i=True,c=GRIS)
d.save(os.path.join(FOLDER,"LIENS_ET_QR_Cousin.docx"))

print("ADMIN CODE:",ADMIN)
print("Fichiers generes dans",FOLDER)
for f in sorted(os.listdir(FOLDER)):print(" -",f)
